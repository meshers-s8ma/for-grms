# app/services/document_service.py

import io
from docx import Document
from docx.text.paragraph import Paragraph

def replace_text_in_paragraph(paragraph: Paragraph, placeholders: dict):
    """
    Находит и заменяет плейсхолдеры в одном параграфе Word-документа.

    Важное замечание: Эта функция объединяет текст из всех "runs" (фрагментов
    текста с разным форматированием) в параграфе, выполняет замену, а затем
    записывает весь измененный текст в первый "run", удаляя остальные.
    Это может привести к потере сложного форматирования внутри параграфа.
    Для простых текстовых замен этот подход работает надежно.

    :param paragraph: Объект параграфа из библиотеки python-docx.
    :param placeholders: Словарь, где ключ - это плейсхолдер (например, '{{Имя}}'),
                         а значение - текст для замены.
    """
    # Собираем весь текст из параграфа воедино
    full_text = "".join(run.text for run in paragraph.runs)

    # Если в тексте нет открывающей скобки, замена не требуется
    if '{' not in full_text:
        return

    # Проходим по всем плейсхолдерам и заменяем их в собранном тексте
    for placeholder, replacement_text in placeholders.items():
        if placeholder in full_text:
            # Убеждаемся, что текст для замены является строкой
            full_text = full_text.replace(placeholder, str(replacement_text))

    # Если в параграфе есть какие-либо "runs"
    if paragraph.runs:
        # Удаляем все "runs", кроме первого, чтобы очистить параграф
        for i in range(len(paragraph.runs) - 1, 0, -1):
            p = paragraph.runs[i]._element
            if p.getparent() is not None:
                p.getparent().remove(p)

        # Записываем весь измененный текст в первый (и теперь единственный) "run"
        paragraph.runs[0].text = full_text


def generate_word_from_data(template_path_or_stream, placeholders: dict) -> io.BytesIO:
    """
    Создает Word-документ на основе шаблона и данных для замены.

    Функция проходит по всем параграфам и таблицам в документе-шаблоне
    и заменяет указанные плейсхолдеры на предоставленные значения.

    :param template_path_or_stream: Путь к файлу шаблона (.docx) или
                                    потоковый объект (например, io.BytesIO).
    :param placeholders: Словарь с данными для замены.
    :return: Потоковый объект io.BytesIO, содержащий сгенерированный Word-документ.
    """
    # Загружаем документ-шаблон из файла или потока
    try:
        doc = Document(template_path_or_stream)
    except Exception as e:
        # Перехватываем возможные ошибки при чтении файла
        raise ValueError(f"Не удалось прочитать шаблон Word. Ошибка: {e}")

    # 1. Замена плейсхолдеров в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, placeholders)

    # 2. Замена плейсхолдеров в основном тексте документа
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)

    # Сохраняем измененный документ в буфер в оперативной памяти
    file_buffer = io.BytesIO()
    doc.save(file_buffer)
    # Перемещаем "курсор" в начало буфера, чтобы его можно было прочитать
    file_buffer.seek(0)

    return file_buffer