# tests/test_services.py

import pytest
import io
import openpyxl
from docx import Document

from app.services import document_service
from app.services import graph_service


class TestDocumentService:
    """Тесты для сервиса генерации Word-документов."""

    def test_generate_word_from_data(self):
        """
        Тест: Проверяет, что плейсхолдеры в шаблоне Word корректно заменяются.
        """
        # 1. Создаем "моковый" шаблон Word в памяти
        doc = Document()
        doc.add_paragraph("Здравствуйте, {{ИМЯ}}!")
        doc.add_paragraph("Добро пожаловать в город {{ГОРОД}}.")
        doc.add_paragraph("Этот текст останется без изменений.")
        
        table = doc.add_table(rows=1, cols=2)
        cell1 = table.cell(0, 0)
        cell1.text = "Ключ: {{КЛЮЧ}}"
        cell2 = table.cell(0, 1)
        cell2.text = "Еще один ключ: {{КЛЮЧ}}"
        
        # Сохраняем шаблон в байтовый поток
        template_stream = io.BytesIO()
        doc.save(template_stream)
        template_stream.seek(0)

        # 2. Определяем данные для замены
        placeholders = {
            "{{ИМЯ}}": "Иван",
            "{{ГОРОД}}": "Москва",
            "{{КЛЮЧ}}": "ЗНАЧЕНИЕ"
        }

        # 3. Вызываем тестируемую функцию
        result_stream = document_service.generate_word_from_data(template_stream, placeholders)

        # 4. Проверяем результат
        result_doc = Document(result_stream)
        
        # Проверяем параграфы
        assert "Здравствуйте, Иван!" in [p.text for p in result_doc.paragraphs]
        assert "Добро пожаловать в город Москва." in [p.text for p in result_doc.paragraphs]
        assert "Этот текст останется без изменений." in [p.text for p in result_doc.paragraphs]
        
        # Проверяем таблицу
        result_table = result_doc.tables[0]
        assert result_table.cell(0, 0).text == "Ключ: ЗНАЧЕНИЕ"
        assert result_table.cell(0, 1).text == "Еще один ключ: ЗНАЧЕНИЕ"


class TestGraphService:
    """Тесты для сервиса работы с Excel-файлами (аналогично Graph API)."""

    def test_read_row_from_excel_bytes_success(self):
        """
        Тест: Проверяет корректное чтение строки из Excel-файла
        и преобразование ее в словарь плейсхолдеров.
        """
        # 1. Создаем "моковый" Excel-файл в памяти
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        
        # Заголовки (проверяем очистку от лишних пробелов)
        sheet["A1"] = " № "
        sheet["B1"] = "Наименование изделия"
        sheet["C1"] = "Количество"
        
        # Строка с данными
        sheet["A2"] = 1
        sheet["B2"] = "Крышка"
        sheet["C2"] = 15.5 # Проверяем, что число преобразуется в строку
        
        excel_stream = io.BytesIO()
        workbook.save(excel_stream)
        excel_bytes = excel_stream.getvalue()

        # 2. Вызываем тестируемую функцию
        placeholders = graph_service.read_row_from_excel_bytes(excel_bytes, row_number=2)

        # 3. Проверяем результат
        expected_placeholders = {
            "{{№}}": "1",
            "{{Наименование изделия}}": "Крышка",
            "{{Количество}}": "15.5"
        }
        assert placeholders == expected_placeholders

    def test_read_row_from_excel_bytes_invalid_row(self):
        """
        Тест: Проверяет, что функция вызывает исключение при неверном номере строки.
        """
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet["A1"] = "Header"
        sheet["A2"] = "Data"
        
        excel_stream = io.BytesIO()
        workbook.save(excel_stream)
        excel_bytes = excel_stream.getvalue()

        # Проверяем, что возникает исключение IndexError, если строка вне диапазона
        with pytest.raises(IndexError):
            graph_service.read_row_from_excel_bytes(excel_bytes, row_number=3)
        
        with pytest.raises(IndexError):
            graph_service.read_row_from_excel_bytes(excel_bytes, row_number=1) # Строка 1 - это заголовки